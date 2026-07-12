"""Publisher 发布器.

输出报告格式需支持 Markdown/HTML/PDF, 默认 Markdown.
"""

from __future__ import annotations

import asyncio
import logging
import re
from typing import Any

from src.config.settings import Settings, get_settings
from src.observability.tracing import trace_chain

logger = logging.getLogger(__name__)


class Publisher:
    """报告发布器.

    支持 Markdown (默认) / HTML / PDF / DOCX / JSON / LaTeX / EPUB 输出
    (扩展 docx + json; 扩展 latex + epub + 多格式同时导出).
    """

    settings: Settings

    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()

    async def publish(
        self,
        report_md: str,
        *,
        output_format: str = "markdown",
        title: str = "",
        sources: list[dict[str, Any]] | None = None,
        agent_role_server: str = "",
        research_mode: str = "",
        user_id: str | None = None,
        session_id: str | None = None,
    ) -> dict[str, Any]:
        """发布报告.

        返回 {"format","content","path"}.
        - markdown: 返回 Markdown 原文
        - html: 返回 HTML 渲染
        - pdf: 返回 PDF 文件路径
        - docx: 返回 DOCX 二进制
        - json: 返回结构化 JSON 字符串
        - latex: 返回 LaTeX 源码 (学术场景)
        - epub: 返回 EPUB 二进制 (电子书场景)
        """
        async with trace_chain(
            name="publisher",
            input={"format": output_format, "report_len": len(report_md)},
            user_id=user_id,
            session_id=session_id,
        ) as span:
            if output_format == "markdown":
                span.update(output={"format": "markdown"})
                return {"format": "markdown", "content": report_md, "path": None}

            if output_format == "html":
                html = await asyncio.to_thread(self._md_to_html, report_md)
                span.update(output={"format": "html", "html_len": len(html)})
                return {"format": "html", "content": html, "path": None}

            if output_format == "pdf":
                pdf_path = await self._md_to_pdf(report_md, session_id)
                span.update(output={"format": "pdf", "pdf_path": pdf_path})
                return {"format": "pdf", "content": None, "path": pdf_path}

            if output_format == "docx":
                docx_bytes = await asyncio.to_thread(self._to_docx, report_md, title=title)
                span.update(output={"format": "docx", "size": len(docx_bytes)})
                return {"format": "docx", "content": docx_bytes, "path": None}

            if output_format == "json":
                json_str = self._to_json(
                    report_md,
                    title=title,
                    sources=sources or [],
                    agent_role_server=agent_role_server,
                    research_mode=research_mode,
                )
                span.update(output={"format": "json", "len": len(json_str)})
                return {"format": "json", "content": json_str, "path": None}

            if output_format == "latex":
                latex = await asyncio.to_thread(self._to_latex, report_md)
                span.update(output={"format": "latex", "len": len(latex)})
                return {"format": "latex", "content": latex, "path": None}

            if output_format == "epub":
                epub_bytes = await asyncio.to_thread(self._to_epub, report_md, title=title)
                span.update(output={"format": "epub", "size": len(epub_bytes)})
                return {"format": "epub", "content": epub_bytes, "path": None}

            # 默认 markdown
            span.update(output={"format": "markdown"})
            return {"format": "markdown", "content": report_md, "path": None}

    def _to_docx(self, content: str, *, title: str = "") -> bytes:
        """Markdown → DOCX (python-docx).

        简易 Markdown 解析 (标题/段落/列表/表格), 失败返回空 bytes.
        中文字体通过修改 Normal 样式 + eastAsia 属性设置, 确保 Word/WPS 正确显示.
        """
        try:
            import io

            from docx import Document
            from docx.oxml.ns import qn
            from docx.shared import Pt

            doc = Document()

            # 设置默认中英文字体 (解决中文乱码)
            # 西文字体用 DejaVu Sans Mono 兼容, 中文用宋体/Noto
            cn_font = "Noto Sans CJK SC"
            en_font = "DejaVu Sans"
            style = doc.styles["Normal"]
            style.font.name = en_font
            style.font.size = Pt(11)
            # eastAsia 属性必须通过 XML 设置, 否则中文显示为方块
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), cn_font)

            if title:
                doc.add_heading(title, level=0)
            # 简易 Markdown 解析 (标题/段落/列表/表格)
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip()
                if not line:
                    i += 1
                    continue
                # SVG 检测 1: <div class="report-image"> 内联 SVG (新格式)
                if '<div class="report-image">' in line:
                    svg_lines: list[str] = []
                    i += 1
                    # 收集 <svg> ... </svg> 内容
                    while i < len(lines) and "</div>" not in lines[i]:
                        svg_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        i += 1  # 跳过 </div>

                    # SVG 转 PNG 后插入 DOCX
                    svg_code = "\n".join(svg_lines)
                    png_bytes = self._svg_to_png_bytes(svg_code)
                    if png_bytes:
                        doc.add_picture(io.BytesIO(png_bytes), width=Pt(400))
                    else:
                        # 转换失败, 降级显示占位文本
                        doc.add_paragraph("[SVG 配图 (转换失败)]")
                    continue
                # SVG 检测 2: ```svg 代码块 (旧格式兼容)
                if line.strip() == "```svg":
                    svg_lines = []
                    i += 1
                    while i < len(lines) and lines[i].strip() != "```":
                        svg_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        i += 1  # 跳过结束 ```

                    # SVG 转 PNG 后插入 DOCX
                    svg_code = "\n".join(svg_lines)
                    png_bytes = self._svg_to_png_bytes(svg_code)
                    if png_bytes:
                        doc.add_picture(io.BytesIO(png_bytes), width=Pt(400))
                    else:
                        # 转换失败, 降级显示占位文本
                        doc.add_paragraph("[SVG 配图 (转换失败)]")
                    continue
                # 表格检测: 当前行以 | 开头, 且下一行是分隔行 (| :--- | :--- |)
                if (
                    line.lstrip().startswith("|")
                    and i + 1 < len(lines)
                    and self._is_md_table_separator(lines[i + 1])
                ):
                    table_lines = [line]
                    i += 1
                    table_lines.append(lines[i])  # 分隔行
                    i += 1
                    while i < len(lines) and lines[i].lstrip().startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    self._add_docx_table(doc, table_lines)
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line[:2].lstrip().isdigit() and ". " in line:
                    # 简单数字列表 (1. xxx / 2. xxx)
                    text = line.split(". ", 1)[-1]
                    doc.add_paragraph(text, style="List Number")
                else:
                    doc.add_paragraph(line)
                i += 1
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e:  # noqa: BLE001
            logger.warning("DOCX 生成失败: %s", e)
            return b""

    @staticmethod
    def _is_md_table_separator(line: str) -> bool:
        """判断一行是否为 Markdown 表格分隔行 (如 | :--- | :--- |)."""
        import re

        stripped = line.strip()
        if not stripped.startswith("|"):
            return False
        inner = stripped.strip("|")
        cells = inner.split("|")
        if not cells:
            return False
        for cell in cells:
            c = cell.strip()
            if not c or not re.fullmatch(r":?-+:?", c):
                return False
        return True

    @staticmethod
    def _parse_md_table_row(line: str) -> list[str]:
        """解析 Markdown 表格行为单元格列表."""
        stripped = line.strip()
        inner = stripped.strip("|")
        return [cell.strip() for cell in inner.split("|")]

    @staticmethod
    def _set_docx_cell_text(cell, text: str, *, bold: bool = False) -> None:
        """设置单元格文本, 支持内联 **bold** 粗体."""
        cell.text = ""
        para = cell.paragraphs[0]
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                run = para.add_run(part)
                if bold:
                    run.bold = True

    def _add_docx_table(self, doc, table_lines: list[str]) -> None:
        """将 Markdown 表格行列表添加到 python-docx Document."""
        header_cells = self._parse_md_table_row(table_lines[0])
        data_rows = [self._parse_md_table_row(line) for line in table_lines[2:]]
        num_cols = len(header_cells)
        if num_cols == 0:
            return
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Table Grid"
        for j, cell_text in enumerate(header_cells):
            if j < num_cols:
                self._set_docx_cell_text(table.rows[0].cells[j], cell_text, bold=True)
        for row_idx, row_cells in enumerate(data_rows):
            for j, cell_text in enumerate(row_cells):
                if j < num_cols:
                    self._set_docx_cell_text(table.rows[row_idx + 1].cells[j], cell_text)

    @staticmethod
    def _svg_to_png_bytes(svg_code: str) -> bytes | None:
        """SVG 代码转 PNG bytes (用 cairosvg).

        cairosvg 不可用时返回 None (降级显示源代码, 不阻断主流程).
        自动从 viewBox 解析原始宽高比, 避免裁剪.
        预处理清理悬空 marker 引用 (cairosvg 2.9.0 bug: 引用缺失时崩溃).
        """
        try:
            import re

            import cairosvg

            # 预处理: 清理悬空 marker 引用 (cairosvg 2.9.0 在 marker-end="url(#id)"
            # 引用的 <marker id="id"> 不存在时抛 AttributeError)
            # 收集所有 def 元素的 id (marker/gradient/filter 等)
            defined_ids = set(re.findall(r'\bid\s*=\s*["\']([^"\']+)["\']', svg_code))
            # 清理 marker-start/marker-mid/marker-end 中引用不存在定义的属性
            def _clean_dangling_marker(m: re.Match) -> str:
                ref_id = m.group(2)
                if ref_id in defined_ids:
                    return m.group(0)
                # 引用不存在, 移除该属性
                return ""
            svg_code = re.sub(
                r'\s(marker-(?:start|mid|end))\s*=\s*["\']url\(#([^)]+)\)["\']',
                _clean_dangling_marker,
                svg_code,
            )

            # 从 SVG viewBox 解析原始宽高比, 避免固定 1024x768 导致裁剪
            # viewBox="0 0 W H" 格式
            viewbox_match = re.search(
                r'viewBox\s*=\s*["\']\s*[\d.]+\s+[\d.]+\s+([\d.]+)\s+([\d.]+)\s*["\']',
                svg_code,
            )
            if viewbox_match:
                vb_w = float(viewbox_match.group(1))
                vb_h = float(viewbox_match.group(2))
                # 以 1024 为基准宽度, 按原始宽高比计算高度
                output_w = 1024
                output_h = int(1024 * vb_h / vb_w)
            else:
                # 无 viewBox, 使用正方形 (避免裁剪)
                output_w = 1024
                output_h = 1024

            return cairosvg.svg2png(
                bytestring=svg_code.encode("utf-8"),
                output_width=output_w,
                output_height=output_h,
            )
        except ImportError:
            logger.warning("cairosvg 未安装, SVG 配图无法转 PNG (DOCX/PDF 降级显示源代码)")
            return None
        except Exception as e:  # noqa: BLE001
            logger.warning("SVG 转 PNG 失败: %s", e)
            return None

    def _to_json(
        self,
        content: str,
        *,
        title: str = "",
        sources: list[dict[str, Any]] | None = None,
        agent_role_server: str = "",
        research_mode: str = "",
    ) -> str:
        """报告 → JSON (结构化输出)."""
        import json
        from datetime import datetime

        report_data = {
            "title": title or "研究报告",
            "content": content,
            "sources": sources or [],
            "metadata": {
                "agent_role_server": agent_role_server,
                "research_mode": research_mode,
                "generated_at": datetime.now().isoformat(),
            },
        }
        return json.dumps(report_data, ensure_ascii=False, indent=2)

    def _md_to_html(self, md: str) -> str:
        """Markdown 转 HTML (用 mistune + jinja2 模板)."""
        try:
            import html

            import mistune

            markdown = mistune.create_markdown(plugins=["table"], escape=False)
            body = markdown(md)

            # 后处理 1: 将 ```svg 代码块转为内联 SVG (浏览器原生渲染)
            # mistune 将 ```svg 渲染为 <pre><code class="language-svg">...</code></pre>
            # 浏览器不直接渲染该格式, 需替换为 <div>{svg_code}</div> 并还原 HTML 实体
            svg_pattern = r'<pre><code class="language-svg">([\s\S]*?)</code></pre>'

            def _unescape_svg(m: re.Match) -> str:
                return f'<div class="report-image">{html.unescape(m.group(1))}</div>'

            body = re.sub(svg_pattern, _unescape_svg, body)

            # 后处理 2: 确保内联 SVG 限制最大宽度 (防止 PDF/HTML 图片超出页面被裁剪)
            # 给 <svg 标签注入 max-width style (如果未有 style 属性)
            # 同时给 div.report-image 添加 text-align:center
            def _constrain_svg(m: re.Match) -> str:
                svg_tag = m.group(0)
                # 如果 <svg 已有 style 属性, 追加 max-width; 否则添加 style 属性
                if 'style="' in svg_tag:
                    svg_tag = svg_tag.replace('style="', 'style="max-width:100%;height:auto;')
                else:
                    svg_tag = svg_tag.replace('<svg ', '<svg style="max-width:100%;height:auto;" ', 1)
                return svg_tag

            body = re.sub(r'<svg[\s]', _constrain_svg, body)

            # HTML 模板 (内联 CSS, 离线友好)
            # 注意: CSS 花括号不需要转义 (用 .replace 不走 Jinja2 模板)
            template = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>研究报告</title>
<style>
body { font-family: "Noto Sans CJK SC", "Noto Sans CJK TC", "Noto Sans CJK JP", "Noto Sans CJK KR", "WenQuanYi Zen Hei", "WenQuanYi Micro Hei", "PingFang SC", "Microsoft YaHei", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; line-height: 1.8; max-width: 800px; margin: 40px auto; padding: 20px; color: #333; }
h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
h2 { color: #34495e; border-bottom: 1px solid #ecf0f1; padding-bottom: 8px; margin-top: 30px; }
h3 { color: #34495e; margin-top: 25px; }
a { color: #3498db; text-decoration: none; }
a:hover { text-decoration: underline; }
code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: "DejaVu Sans Mono", "Source Code Pro", monospace; }
blockquote { border-left: 4px solid #3498db; margin: 0; padding-left: 16px; color: #666; }
table { border-collapse: collapse; width: 100%; margin: 16px 0; }
th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; }
th { background: #f8f9fa; font-weight: 600; }
.report-image { text-align: center; margin: 20px 0; }
.report-image svg { max-width: 100%; height: auto; }
</style>
</head>
<body>
__BODY_PLACEHOLDER__
</body>
</html>"""
            # 用 placeholder 避免 {body} 与 CSS 中的 { } 冲突
            return template.replace("__BODY_PLACEHOLDER__", body)
        except Exception as e:  # noqa: BLE001
            logger.warning("HTML 渲染失败, 返回纯文本: %s", e)
            return f"<html><body><pre>{md}</pre></body></html>"

    async def _md_to_pdf(self, md: str, session_id: str | None = None) -> str:
        """Markdown 转 PDF (WeasyPrint)."""
        try:
            import asyncio
            import os

            # 先转 HTML
            html = self._md_to_html(md)

            # PDF 输出路径
            upload_dir = self.settings.upload_dir
            os.makedirs(upload_dir, exist_ok=True)
            filename = f"report_{session_id or 'unknown'}.pdf"
            pdf_path = os.path.join(upload_dir, filename)

            def _sync_pdf() -> None:
                from weasyprint import HTML

                HTML(string=html).write_pdf(pdf_path)

            await asyncio.to_thread(_sync_pdf)
            return pdf_path
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF 生成失败: %s", e)
            # 降级返回 HTML
            html_path = await self._save_html_fallback(md, session_id)
            return html_path

    async def _save_html_fallback(self, md: str, session_id: str | None = None) -> str:
        """PDF 失败时降级保存 HTML."""
        import asyncio
        import os

        upload_dir = self.settings.upload_dir
        os.makedirs(upload_dir, exist_ok=True)
        filename = f"report_{session_id or 'unknown'}.html"
        html_path = os.path.join(upload_dir, filename)
        html = self._md_to_html(md)

        def _write() -> None:
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html)

        await asyncio.to_thread(_write)
        return html_path

    def _to_latex(self, report_md: str) -> str:
        """Markdown → LaTeX (学术场景).

        基础转换: 标题/列表/粗体, 纯 Python 实现不引入新依赖.
        """
        import re

        lines = report_md.split("\n")
        latex_lines: list[str] = [
            "\\documentclass[12pt]{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{hyperref}",
            "\\title{Research Report}",
            "\\date{}",
            "\\begin{document}",
            "\\maketitle",
        ]

        body_lines: list[str] = []
        in_list = False
        for line in lines:
            if line.startswith("# "):
                if in_list:
                    body_lines.append("\\end{itemize}")
                    in_list = False
                body_lines.append(f"\\section{{{line[2:]}}}")
            elif line.startswith("## "):
                if in_list:
                    body_lines.append("\\end{itemize}")
                    in_list = False
                body_lines.append(f"\\subsection{{{line[3:]}}}")
            elif line.startswith("### "):
                if in_list:
                    body_lines.append("\\end{itemize}")
                    in_list = False
                body_lines.append(f"\\subsubsection{{{line[4:]}}}")
            elif line.startswith("- ") or line.startswith("* "):
                if not in_list:
                    body_lines.append("\\begin{itemize}")
                    in_list = True
                body_lines.append(f"\\item {line[2:]}")
            elif line.strip():
                if in_list:
                    body_lines.append("\\end{itemize}")
                    in_list = False
                # 粗体转换 **bold** → \textbf{bold}
                converted = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", line)
                body_lines.append(converted)
            else:
                if in_list:
                    body_lines.append("\\end{itemize}")
                    in_list = False
                body_lines.append("")

        if in_list:
            body_lines.append("\\end{itemize}")

        latex_lines.extend(body_lines)
        latex_lines.append("\\end{document}")
        return "\n".join(latex_lines)

    def _to_epub(self, report_md: str, *, title: str = "") -> bytes:
        """Markdown → EPUB (电子书场景, 纯 stdlib zipfile 实现)."""
        import io
        import re
        import uuid
        import zipfile
        from datetime import datetime
        from xml.sax.saxutils import escape

        try:
            html_body = self._md_to_html(report_md)
            # _md_to_html 返回完整 HTML 文档, 提取 <body>...</body>
            body_match = re.search(r"<body>(.*)</body>", html_body, re.DOTALL)
            body_content = body_match.group(1) if body_match else html_body

            book_title = title or "研究报告"
            author = "AgentInsight Researcher"
            language = "zh-CN"
            identifier = f"urn:uuid:{uuid.uuid4()}"
            created = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

            container_xml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<container version="1.0" '
                'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
                "  <rootfiles>\n"
                '    <rootfile full-path="OEBPS/content.opf" '
                'media-type="application/oebps-package+xml"/>\n'
                "  </rootfiles>\n"
                "</container>"
            )

            content_opf = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" '
                'unique-identifier="BookId">\n'
                '  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" '
                'xmlns:opf="http://www.idpf.org/2007/opf">\n'
                f"    <dc:title>{escape(book_title)}</dc:title>\n"
                f"    <dc:creator>{escape(author)}</dc:creator>\n"
                f"    <dc:language>{language}</dc:language>\n"
                f'    <dc:identifier id="BookId">{identifier}</dc:identifier>\n'
                f"    <dc:date>{created}</dc:date>\n"
                "  </metadata>\n"
                "  <manifest>\n"
                '    <item id="content" href="content.xhtml" '
                'media-type="application/xhtml+xml"/>\n'
                '    <item id="ncx" href="toc.ncx" '
                'media-type="application/x-dtbncx+xml"/>\n'
                "  </manifest>\n"
                '  <spine toc="ncx">\n'
                '    <itemref idref="content"/>\n'
                "  </spine>\n"
                "</package>"
            )

            toc_ncx = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">\n'
                "  <head>\n"
                f'    <meta name="dtb:uid" content="{identifier}"/>\n'
                "  </head>\n"
                f"  <docTitle><text>{escape(book_title)}</text></docTitle>\n"
                "  <navMap>\n"
                '    <navPoint id="navpoint-1" playOrder="1">\n'
                f"      <navLabel><text>{escape(book_title)}</text></navLabel>\n"
                '      <content src="content.xhtml"/>\n'
                "    </navPoint>\n"
                "  </navMap>\n"
                "</ncx>"
            )

            xhtml = (
                '<?xml version="1.0" encoding="UTF-8"?>\n'
                '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" '
                '"http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">\n'
                '<html xmlns="http://www.w3.org/1999/xhtml">\n'
                "<head>\n"
                f"<title>{escape(book_title)}</title>\n"
                "</head>\n"
                "<body>\n"
                f"{body_content}\n"
                "</body>\n"
                "</html>"
            )

            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                # mimetype 必须第一个且不压缩 (EPUB 规范)
                zf.writestr(
                    "mimetype",
                    "application/epub+zip",
                    compress_type=zipfile.ZIP_STORED,
                )
                zf.writestr("META-INF/container.xml", container_xml)
                zf.writestr("OEBPS/content.opf", content_opf)
                zf.writestr("OEBPS/toc.ncx", toc_ncx)
                zf.writestr("OEBPS/content.xhtml", xhtml)
            return buf.getvalue()
        except Exception as e:  # noqa: BLE001
            logger.warning("EPUB 生成失败: %s", e)
            return b""

    async def export_multiple_formats(
        self,
        report_md: str,
        formats: list[str],
        *,
        title: str = "",
        sources: list[dict[str, Any]] | None = None,
        agent_role_server: str = "",
        research_mode: str = "",
        user_id: str | None = None,
        session_id: str | None = None,
    ) -> dict[str, Any]:
        """一次报告生成多种格式.

        返回 dict, key 为格式名 (markdown/html/pdf_path/docx/json/latex/epub),
        value 为对应内容 (字符串/字节/文件路径).

        改用 asyncio.gather 并行执行多种格式导出, 延迟不再随格式数线性增长;
        return_exceptions=True 隔离单个格式失败, 不影响其他格式.
        每个 publish() 内部已有 trace_chain 包裹, 并行后仍保留追踪能力.

        接入主流程: routes.py 的 download_report 端点可接收 formats=list[str] 参数,
        调用此方法批量导出. user_id/session_id 透传给 publish() 用于 trace_chain.
        """
        # 格式名 → 结果 key 映射; pdf 取 path 字段, 其他取 content 字段
        fmt_to_key = {
            "markdown": "markdown",
            "html": "html",
            "pdf": "pdf_path",
            "docx": "docx",
            "json": "json",
            "latex": "latex",
            "epub": "epub",
        }

        # 过滤未知格式 (与原逻辑一致, 跳过并 warning)
        valid_formats: list[str] = []
        for fmt in formats:
            fmt_lower = fmt.lower()
            if fmt_lower in fmt_to_key:
                valid_formats.append(fmt_lower)
            else:
                logger.warning("未知导出格式: %s (跳过)", fmt)

        # 并行调用 publish(), 每个调用内部 trace_chain 包裹保留
        gathered = await asyncio.gather(
            *[
                self.publish(
                    report_md,
                    output_format=fmt,
                    title=title,
                    sources=sources,
                    agent_role_server=agent_role_server,
                    research_mode=research_mode,
                    user_id=user_id,
                    session_id=session_id,
                )
                for fmt in valid_formats
            ],
            return_exceptions=True,
        )

        # 按 key 组织结果; 单格式异常隔离, 仅 warning 不影响其他格式
        results: dict[str, Any] = {}
        for fmt, outcome in zip(valid_formats, gathered, strict=False):
            if isinstance(outcome, BaseException):
                logger.warning("格式 %s 导出失败: %s", fmt, outcome)
                continue
            key = fmt_to_key[fmt]
            # pdf 格式返回文件路径, 其他格式返回 content
            results[key] = outcome.get("path") if fmt == "pdf" else outcome.get("content")
        return results
