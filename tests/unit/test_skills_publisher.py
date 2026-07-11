"""单元测试: Publisher 发布器纯转换函数.

验证 Markdown → LaTeX / EPUB / JSON / DOCX / HTML 转换逻辑,
不依赖外部服务 (LLM / 文件系统).
单元测试在构建期执行, 不依赖外部服务.
"""

from __future__ import annotations

import json
import zipfile

import pytest

from src.config.settings import Settings
from src.skills.researcher.publisher import Publisher

pytestmark = pytest.mark.unit


@pytest.fixture()
def publisher() -> Publisher:
    """构造 Publisher 实例 (隔离 .env)."""
    return Publisher(settings=Settings(_env_file=None))


# ========== _to_latex ==========


def test_to_latex_headings(publisher: Publisher) -> None:
    """测试 Markdown 标题 (# / ## / ###) 转 LaTeX section."""
    md = "# 一级标题\n## 二级标题\n### 三级标题"
    latex = publisher._to_latex(md)
    assert "\\section{一级标题}" in latex
    assert "\\subsection{二级标题}" in latex
    assert "\\subsubsection{三级标题}" in latex
    assert "\\documentclass[12pt]{article}" in latex
    assert "\\begin{document}" in latex
    assert "\\end{document}" in latex


def test_to_latex_unordered_list(publisher: Publisher) -> None:
    """测试 Markdown 无序列表 (- / *) 转 LaTeX itemize."""
    md = "- 项目一\n- 项目二\n* 项目三"
    latex = publisher._to_latex(md)
    assert "\\begin{itemize}" in latex
    assert "\\end{itemize}" in latex
    assert "\\item 项目一" in latex
    assert "\\item 项目二" in latex
    assert "\\item 项目三" in latex


def test_to_latex_bold(publisher: Publisher) -> None:
    """测试 Markdown 粗体 (**text**) 转 LaTeX textbf."""
    md = "这是一段**重点强调**的正文"
    latex = publisher._to_latex(md)
    assert "\\textbf{重点强调}" in latex
    assert "**" not in latex


def test_to_latex_plain_paragraph(publisher: Publisher) -> None:
    """测试普通段落保留为 LaTeX 纯文本."""
    md = "这是一段普通正文内容"
    latex = publisher._to_latex(md)
    assert "这是一段普通正文内容" in latex


def test_to_latex_list_closes_before_heading(publisher: Publisher) -> None:
    """测试列表遇到新标题时正确关闭 itemize."""
    md = "- 项目一\n# 新标题"
    latex = publisher._to_latex(md)
    # itemize 必须在新 section 之前关闭
    section_pos = latex.find("\\section{新标题}")
    end_itemize_pos = latex.find("\\end{itemize}")
    assert end_itemize_pos < section_pos
    assert end_itemize_pos != -1


# ========== _to_epub ==========


def test_to_epub_returns_bytes(publisher: Publisher) -> None:
    """测试 EPUB 返回字节对象."""
    epub = publisher._to_epub("# 测试报告\n正文内容", title="测试标题")
    assert isinstance(epub, bytes)
    assert len(epub) > 0


def test_to_epub_zip_structure(publisher: Publisher) -> None:
    """测试 EPUB ZIP 含 mimetype/container.opf/toc.ncx/content.xhtml."""
    epub = publisher._to_epub("# 测试报告", title="测试标题")
    with zipfile.ZipFile(io_from_bytes(epub)) as zf:
        names = zf.namelist()
        assert "mimetype" in names
        assert "META-INF/container.xml" in names
        assert "OEBPS/content.opf" in names
        assert "OEBPS/toc.ncx" in names
        assert "OEBPS/content.xhtml" in names


def test_to_epub_contains_title(publisher: Publisher) -> None:
    """测试 EPUB 元数据含 title."""
    epub = publisher._to_epub("# 报告内容", title="我的研究报告")
    with zipfile.ZipFile(io_from_bytes(epub)) as zf:
        opf = zf.read("OEBPS/content.opf").decode("utf-8")
        ncx = zf.read("OEBPS/toc.ncx").decode("utf-8")
        xhtml = zf.read("OEBPS/content.xhtml").decode("utf-8")
    assert "我的研究报告" in opf
    assert "我的研究报告" in ncx
    assert "我的研究报告" in xhtml


def test_to_epub_mimetype_first_and_uncompressed(publisher: Publisher) -> None:
    """测试 mimetype 为 EPUB 第一个文件且不压缩 (规范要求)."""
    epub = publisher._to_epub("# x", title="t")
    with zipfile.ZipFile(io_from_bytes(epub)) as zf:
        info = zf.infolist()[0]
        assert info.filename == "mimetype"
        assert info.compress_type == zipfile.ZIP_STORED


def test_to_epub_default_title(publisher: Publisher) -> None:
    """测试未传 title 时使用默认 '研究报告'."""
    epub = publisher._to_epub("# 内容")
    with zipfile.ZipFile(io_from_bytes(epub)) as zf:
        opf = zf.read("OEBPS/content.opf").decode("utf-8")
    assert "研究报告" in opf


# ========== _to_json ==========


def test_to_json_structure(publisher: Publisher) -> None:
    """测试 JSON 输出含 title/content/sources/metadata 四键."""
    sources = [{"title": "src1", "url": "https://example.com"}]
    json_str = publisher._to_json(
        "报告内容",
        title="我的报告",
        sources=sources,
        agent_role_server="analyst",
        research_mode="basic",
    )
    data = json.loads(json_str)
    assert data["title"] == "我的报告"
    assert data["content"] == "报告内容"
    assert data["sources"] == sources
    assert data["metadata"]["agent_role_server"] == "analyst"
    assert data["metadata"]["research_mode"] == "basic"
    assert "generated_at" in data["metadata"]


def test_to_json_default_title(publisher: Publisher) -> None:
    """测试未传 title 时默认 '研究报告'."""
    json_str = publisher._to_json("内容")
    data = json.loads(json_str)
    assert data["title"] == "研究报告"
    assert data["sources"] == []
    assert data["metadata"]["agent_role_server"] == ""
    assert data["metadata"]["research_mode"] == ""


def test_to_json_is_serializable(publisher: Publisher) -> None:
    """测试 _to_json 返回的是合法 JSON 字符串."""
    json_str = publisher._to_json("内容", title="t")
    assert isinstance(json_str, str)
    # 不应抛异常
    json.loads(json_str)


# ========== _to_docx ==========


def test_to_docx_returns_bytes(publisher: Publisher) -> None:
    """测试 DOCX 返回字节对象."""
    docx_bytes = publisher._to_docx("# 标题\n正文", title="测试")
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0


def test_to_docx_failure_returns_empty(publisher: Publisher) -> None:
    """测试 DOCX 生成失败时返回空 bytes (不抛异常)."""
    # 用 monkeypatch 模拟 import 失败
    import sys

    original_docx = sys.modules.get("docx")
    sys.modules["docx"] = None  # type: ignore[assignment]
    try:
        result = publisher._to_docx("# 标题", title="t")
        assert result == b""
    finally:
        if original_docx is not None:
            sys.modules["docx"] = original_docx
        else:
            sys.modules.pop("docx", None)


# ========== _md_to_html ==========


def test_md_to_html_contains_inline_css(publisher: Publisher) -> None:
    """测试 HTML 输出含内联 CSS (<style> 标签)."""
    html = publisher._md_to_html("# 标题\n正文")
    assert "<style>" in html
    assert "font-family" in html
    assert "<!DOCTYPE html>" in html


def test_md_to_html_renders_markdown(publisher: Publisher) -> None:
    """测试 mistune 渲染 Markdown (H1 → <h1>)."""
    html = publisher._md_to_html("# 我的标题\n\n一段正文")
    assert "<h1>" in html
    assert "我的标题" in html
    assert "<p>" in html


def test_md_to_html_renders_list(publisher: Publisher) -> None:
    """测试 mistune 渲染列表."""
    html = publisher._md_to_html("- 项一\n- 项二")
    assert "<ul>" in html
    assert "<li>" in html


def test_md_to_html_fallback_without_mistune(publisher: Publisher) -> None:
    """测试 mistune 未安装时降级为 <pre> 包裹的纯文本 (不抛异常).

    用 sys.modules mock 模拟 mistune 不可用, 确保测试不依赖环境.
    """
    import sys

    original_mistune = sys.modules.get("mistune")
    sys.modules["mistune"] = None  # type: ignore[assignment]
    try:
        html = publisher._md_to_html("# 标题\n正文")
        assert "<html>" in html
        assert "<pre>" in html
        assert "# 标题" in html
    finally:
        if original_mistune is not None:
            sys.modules["mistune"] = original_mistune
        else:
            sys.modules.pop("mistune", None)


def test_md_to_html_renders_table(publisher: Publisher) -> None:
    """测试 mistune table 插件渲染 Markdown 表格为 <table>."""
    md = "| 列A | 列B |\n| :--- | :--- |\n| 1 | 2 |\n| 3 | 4 |"
    html = publisher._md_to_html(md)
    assert "<table>" in html
    assert "<thead>" in html
    assert "<tbody>" in html
    assert "</td>" in html
    assert "列A" in html
    assert "列B" in html


def test_md_to_html_table_not_plain_text(publisher: Publisher) -> None:
    """测试表格不被渲染为纯文本段落 (回归测试)."""
    md = "| 名称 | 数值 |\n| :--- | :--- |\n| alpha | 100 |"
    html = publisher._md_to_html(md)
    # 不应出现把 | 当作普通文本的 <p>| 名称...</p>
    assert "<table>" in html
    assert "<p>| 名称" not in html


# ========== _to_docx 表格 ==========


def test_to_docx_contains_table(publisher: Publisher) -> None:
    """测试 DOCX 输出包含表格对象."""
    import io

    from docx import Document

    md = "# 标题\n\n| 列A | 列B |\n| :--- | :--- |\n| 1 | 2 |\n"
    docx_bytes = publisher._to_docx(md, title="测试")
    assert isinstance(docx_bytes, bytes)
    doc = Document(io.BytesIO(docx_bytes))
    tables = doc.tables
    assert len(tables) == 1
    assert len(tables[0].rows) == 2  # header + 1 data row
    assert len(tables[0].columns) == 2
    assert tables[0].rows[0].cells[0].text == "列A"
    assert tables[0].rows[0].cells[1].text == "列B"
    assert tables[0].rows[1].cells[0].text == "1"
    assert tables[0].rows[1].cells[1].text == "2"


def test_to_docx_table_bold_in_cell(publisher: Publisher) -> None:
    """测试 DOCX 表格单元格内联 **bold** 粗体."""
    import io

    from docx import Document

    md = "| 名称 | 值 |\n| :--- | :--- |\n| **重点** | 100 |\n"
    docx_bytes = publisher._to_docx(md, title="t")
    doc = Document(io.BytesIO(docx_bytes))
    cell = doc.tables[0].rows[1].cells[0]
    # cell.text 去除 markdown 标记
    assert cell.text == "重点"
    # 至少一个 run 应为 bold
    assert any(r.bold for p in cell.paragraphs for r in (p.runs or []))


def test_is_md_table_separator(publisher: Publisher) -> None:
    """测试表格分隔行检测."""
    assert publisher._is_md_table_separator("| :--- | :--- |")
    assert publisher._is_md_table_separator("| --- | --- |")
    assert publisher._is_md_table_separator("|:---|:---|")
    assert publisher._is_md_table_separator("| :---: | ---: | :--- |")
    assert not publisher._is_md_table_separator("| 名称 | 值 |")
    assert not publisher._is_md_table_separator("普通文本")
    assert not publisher._is_md_table_separator("| abc | def |")


# ========== 辅助函数 ==========


def io_from_bytes(data: bytes):
    """BytesIO 包装工具 (避免顶层 import io)."""
    import io

    return io.BytesIO(data)
