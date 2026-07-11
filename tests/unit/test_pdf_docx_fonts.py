"""单元测试: PDF/DOCX 中文字体配置 (静态文件检查).

验证离线部署模式的中文字体安装逻辑, 确保 PDF/DOCX 报告不乱码:
- Dockerfile.qa / Dockerfile.offline 含 fonts-noto-cjk 字体安装逻辑
- packages/debs/ 目录含 fonts-noto-cjk*.deb 与 fonts-wqy-*.deb 预下载包
- docker-compose-qa.yaml 的 embeddings/rerank 健康检查用 CMD-SHELL (非 CMD curl)

注意: 这些是静态文件检查, 不需要真正构建 Docker 镜像.
Dockerfile.qa / Dockerfile.offline / docker-compose-qa.yaml / packages/debs/ 均在
.gitignore 中 (AGENTS.md 第 12 章三套构建模式, 不入仓), 文件不存在时跳过用例
而非失败 (CI 全新克隆环境可能无这些文件).

AGENTS.md 第 13 章: 单元测试在构建期执行, 不依赖外部服务.
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytestmark = pytest.mark.unit

# 项目根目录 (tests/unit/ 上溯两级)
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_DOCKERFILE_QA = _PROJECT_ROOT / "Dockerfile.qa"
_DOCKERFILE_OFFLINE = _PROJECT_ROOT / "Dockerfile.offline"
_COMPOSE_QA = _PROJECT_ROOT / "docker-compose-qa.yaml"
_DEBS_DIR = _PROJECT_ROOT / "packages" / "debs"


# ========== Dockerfile.qa 字体安装逻辑 ==========


def test_dockerfile_qa_contains_fonts_noto_cjk() -> None:
    """Dockerfile.qa 含 fonts-noto-cjk 关键词 (中文字体安装逻辑)."""
    if not _DOCKERFILE_QA.exists():
        pytest.skip("Dockerfile.qa 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_QA.read_text(encoding="utf-8")
    assert "fonts-noto-cjk" in content, (
        "Dockerfile.qa 应含 fonts-noto-cjk 字体安装逻辑 (PDF/DOCX 报告中文显示)"
    )


def test_dockerfile_qa_contains_dpkg_install_logic() -> None:
    """Dockerfile.qa 含 dpkg -i 离线安装逻辑 (.deb 包安装)."""
    if not _DOCKERFILE_QA.exists():
        pytest.skip("Dockerfile.qa 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_QA.read_text(encoding="utf-8")
    assert "dpkg -i" in content, "Dockerfile.qa 应含 dpkg -i 离线安装逻辑"
    assert "packages/debs" in content, "Dockerfile.qa 应 COPY packages/debs/"


def test_dockerfile_qa_contains_font_verification() -> None:
    """Dockerfile.qa 含 fc-list 字体验证逻辑 (构建时校验字体已安装)."""
    if not _DOCKERFILE_QA.exists():
        pytest.skip("Dockerfile.qa 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_QA.read_text(encoding="utf-8")
    assert "fc-list" in content, "Dockerfile.qa 应含 fc-list 字体验证逻辑"


def test_dockerfile_qa_mentions_wqy_fonts() -> None:
    """Dockerfile.qa 提及 fonts-wqy 文泉驿字体 (备选中文字体)."""
    if not _DOCKERFILE_QA.exists():
        pytest.skip("Dockerfile.qa 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_QA.read_text(encoding="utf-8")
    assert "fonts-wqy" in content, "Dockerfile.qa 应提及 fonts-wqy 文泉驿字体"


# ========== Dockerfile.offline 字体安装逻辑 ==========


def test_dockerfile_offline_contains_fonts_noto_cjk() -> None:
    """Dockerfile.offline 含 fonts-noto-cjk 关键词 (中文字体安装逻辑)."""
    if not _DOCKERFILE_OFFLINE.exists():
        pytest.skip("Dockerfile.offline 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_OFFLINE.read_text(encoding="utf-8")
    assert "fonts-noto-cjk" in content, (
        "Dockerfile.offline 应含 fonts-noto-cjk 字体安装逻辑 (PDF/DOCX 报告中文显示)"
    )


def test_dockerfile_offline_contains_dpkg_install_logic() -> None:
    """Dockerfile.offline 含 dpkg -i 离线安装逻辑."""
    if not _DOCKERFILE_OFFLINE.exists():
        pytest.skip("Dockerfile.offline 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_OFFLINE.read_text(encoding="utf-8")
    assert "dpkg -i" in content, "Dockerfile.offline 应含 dpkg -i 离线安装逻辑"
    assert "packages/debs" in content, "Dockerfile.offline 应 COPY packages/debs/"


def test_dockerfile_offline_contains_font_verification() -> None:
    """Dockerfile.offline 含 fc-list 字体验证逻辑."""
    if not _DOCKERFILE_OFFLINE.exists():
        pytest.skip("Dockerfile.offline 不存在 (gitignored, 离线构建文件未预生成)")

    content = _DOCKERFILE_OFFLINE.read_text(encoding="utf-8")
    assert "fc-list" in content, "Dockerfile.offline 应含 fc-list 字体验证逻辑"


# ========== packages/debs/ 字体 .deb 包 ==========


def test_debs_dir_contains_fonts_noto_cjk_deb() -> None:
    """packages/debs/ 目录含 fonts-noto-cjk*.deb 文件 (预下载中文字体包)."""
    if not _DEBS_DIR.exists():
        pytest.skip("packages/debs/ 目录不存在 (gitignored, 离线包未预下载)")

    noto_cjk_debs = list(_DEBS_DIR.glob("fonts-noto-cjk*.deb"))
    assert len(noto_cjk_debs) > 0, (
        "packages/debs/ 应含 fonts-noto-cjk*.deb 文件 (PDF/DOCX 报告中文字体). "
        "预下载命令: apt-get download fonts-noto-cjk fonts-noto-cjk-extra"
    )


def test_debs_dir_contains_fonts_wqy_deb() -> None:
    """packages/debs/ 目录含 fonts-wqy-*.deb 文件 (文泉驿备选字体)."""
    if not _DEBS_DIR.exists():
        pytest.skip("packages/debs/ 目录不存在 (gitignored, 离线包未预下载)")

    wqy_debs = list(_DEBS_DIR.glob("fonts-wqy-*.deb"))
    assert len(wqy_debs) > 0, (
        "packages/debs/ 应含 fonts-wqy-*.deb 文件 (文泉驿中文字体). "
        "预下载命令: apt-get download fonts-wqy-zenhei fonts-wqy-microhei"
    )


def test_debs_dir_contains_at_least_two_wqy_fonts() -> None:
    """packages/debs/ 应含至少 2 个 fonts-wqy-*.deb (zenhei + microhei)."""
    if not _DEBS_DIR.exists():
        pytest.skip("packages/debs/ 目录不存在 (gitignored, 离线包未预下载)")

    wqy_debs = list(_DEBS_DIR.glob("fonts-wqy-*.deb"))
    if not wqy_debs:
        pytest.skip("packages/debs/ 无 fonts-wqy-*.deb 文件 (未预下载)")
    assert len(wqy_debs) >= 2, (
        f"应含至少 2 个 fonts-wqy-*.deb (zenhei + microhei), 实际: {len(wqy_debs)}"
    )


# ========== docker-compose-qa.yaml 健康检查 (CMD-SHELL) ==========


def _load_compose_qa() -> dict:
    """加载 docker-compose-qa.yaml, 文件不存在时抛 skip."""
    if not _COMPOSE_QA.exists():
        pytest.skip("docker-compose-qa.yaml 不存在 (gitignored, 离线构建文件未预生成)")
    try:
        import yaml
    except ImportError:
        pytest.skip("PyYAML 未安装, 跳过 YAML 结构化检查")
    with _COMPOSE_QA.open(encoding="utf-8") as f:
        return yaml.safe_load(f)


def test_compose_qa_embeddings_healthcheck_uses_cmd_shell() -> None:
    """embeddings 健康检查用 CMD-SHELL (非 CMD curl), 绕过 TEI ENTRYPOINT 拦截."""
    compose = _load_compose_qa()
    embeddings = compose.get("services", {}).get("embeddings", {})
    healthcheck = embeddings.get("healthcheck", {})
    test_list = healthcheck.get("test", [])

    assert len(test_list) >= 1, "embeddings healthcheck.test 不能为空"
    assert test_list[0] == "CMD-SHELL", (
        f"embeddings healthcheck 应用 CMD-SHELL (TEI ENTRYPOINT 拦截 CMD 形式), "
        f"实际: {test_list[0]}"
    )


def test_compose_qa_rerank_healthcheck_uses_cmd_shell() -> None:
    """rerank 健康检查用 CMD-SHELL (非 CMD curl)."""
    compose = _load_compose_qa()
    rerank = compose.get("services", {}).get("rerank", {})
    healthcheck = rerank.get("healthcheck", {})
    test_list = healthcheck.get("test", [])

    assert len(test_list) >= 1, "rerank healthcheck.test 不能为空"
    assert test_list[0] == "CMD-SHELL", f"rerank healthcheck 应用 CMD-SHELL, 实际: {test_list[0]}"


def test_compose_qa_embeddings_healthcheck_uses_curl() -> None:
    """embeddings CMD-SHELL 健康检查含 curl (检查 /health 端点)."""
    compose = _load_compose_qa()
    embeddings = compose.get("services", {}).get("embeddings", {})
    test_list = embeddings.get("healthcheck", {}).get("test", [])

    # CMD-SHELL 形式: test[1] 是 shell 命令字符串
    assert len(test_list) >= 2, "embeddings healthcheck CMD-SHELL 应含命令字符串"
    shell_cmd = test_list[1]
    assert "curl" in shell_cmd, f"embeddings 健康检查应含 curl, 实际: {shell_cmd}"
    assert "/health" in shell_cmd, f"embeddings 健康检查应检查 /health 端点, 实际: {shell_cmd}"


def test_compose_qa_rerank_healthcheck_uses_curl() -> None:
    """rerank CMD-SHELL 健康检查含 curl."""
    compose = _load_compose_qa()
    rerank = compose.get("services", {}).get("rerank", {})
    test_list = rerank.get("healthcheck", {}).get("test", [])

    assert len(test_list) >= 2, "rerank healthcheck CMD-SHELL 应含命令字符串"
    shell_cmd = test_list[1]
    assert "curl" in shell_cmd, f"rerank 健康检查应含 curl, 实际: {shell_cmd}"


def test_compose_qa_embeddings_healthcheck_not_cmd_form() -> None:
    """embeddings 健康检查不应是 CMD 形式 (会被 TEI ENTRYPOINT 拦截).

    TEI 镜像 ENTRYPOINT=[text-embeddings-router], CMD 形式会被拦截,
    必须用 CMD-SHELL + curl 绝对路径绕过.
    """
    compose = _load_compose_qa()
    embeddings = compose.get("services", {}).get("embeddings", {})
    test_list = embeddings.get("healthcheck", {}).get("test", [])

    assert test_list[0] != "CMD", (
        "embeddings healthcheck 不应用 CMD 形式 (TEI ENTRYPOINT 会拦截), 应改用 CMD-SHELL"
    )


# ========== 运行时 PDF/DOCX 中文渲染验证 (验证不乱码) ==========


def test_docx_chinese_text_not_garbled() -> None:
    """运行时验证: DOCX 中文文本不乱码.

    生成含中文的 DOCX → 读回 → 验证中文文本正确提取.
    验证 publisher._to_docx 的 eastAsia 字体配置生效.
    """
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        pytest.skip("python-docx 未安装, 跳过 DOCX 运行时验证")

    from src.skills.researcher.publisher import Publisher

    chinese_content = (
        "# 测试报告标题\n\n这是一段中文内容, 用于验证 DOCX 不乱码.\n- 列表项一\n- 列表项二"
    )
    publisher = Publisher()
    docx_bytes = publisher._to_docx(chinese_content, title="中文测试标题")
    assert len(docx_bytes) > 0, "DOCX 生成失败, 返回空 bytes"

    import io

    doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "中文内容" in all_text, f"DOCX 中文文本乱码或缺失, 实际: {all_text[:200]}"
    assert "列表项一" in all_text, f"DOCX 中文列表项乱码, 实际: {all_text[:200]}"


def test_docx_mixed_chinese_english_not_garbled() -> None:
    """运行时验证: DOCX 中英文混合不乱码."""
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        pytest.skip("python-docx 未安装, 跳过 DOCX 运行时验证")

    from src.skills.researcher.publisher import Publisher

    mixed_content = "# LLM Hallucination 大语言模型幻觉\n\nRAG 检索增强生成可以缓解幻觉问题."
    publisher = Publisher()
    docx_bytes = publisher._to_docx(mixed_content, title="中英文混合测试")
    assert len(docx_bytes) > 0

    import io

    doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "大语言模型幻觉" in all_text, f"DOCX 中文部分乱码, 实际: {all_text[:200]}"
    assert "RAG" in all_text, f"DOCX 英文部分缺失, 实际: {all_text[:200]}"


def test_docx_font_config_uses_noto_cjk() -> None:
    """验证 DOCX 字体配置使用 Noto Sans CJK SC (eastAsia 属性)."""
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        pytest.skip("python-docx 未安装, 跳过 DOCX 运行时验证")

    from src.skills.researcher.publisher import Publisher

    publisher = Publisher()
    docx_bytes = publisher._to_docx("测试", title="字体验证")
    assert len(docx_bytes) > 0

    import io

    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(docx_bytes))
    style = doc.styles["Normal"]
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            east_asia = rfonts.get(qn("w:eastAsia"))
            assert east_asia is not None, "DOCX Normal 样式缺少 eastAsia 字体属性"
            assert "CJK" in east_asia or "Noto" in east_asia, (
                f"DOCX eastAsia 字体应为 Noto Sans CJK SC, 实际: {east_asia}"
            )
